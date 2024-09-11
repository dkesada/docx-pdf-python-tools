from docx import Document
import sys
import os
import comtypes.client
from docx2pdf import convert
from tqdm import tqdm


def get_para_data(output_doc_name, paragraph):
    # Adapted from https://stackoverflow.com/a/65956636
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        output_run.font.size = run.font.size
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    output_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    output_para.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
    output_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
    output_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
    output_para.paragraph_format.space_before = paragraph.paragraph_format.space_before


def save_file(name):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(name + '.docx')
    doc.SaveAs(name + '.pdf', FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


# Convert each page of a docx into an independent pdf file

wdFormatPDF = 17
file_name = sys.argv[1]
doc = Document(file_name)
new_doc = Document()

# To differentiate pages, one needs to check the beginning or ending of paragraphs
new_page = 'Name:'  # If all pages start with "Name:" followed by a name, for example
name = doc.paragraphs[0].text.split(': ')[1]
i = 0

for par in tqdm(doc.paragraphs[1:]):
    if par.text.startswith(new_page):  # New page, create new doc
        new_doc.save(name + '.docx')
        save_file(name)
        name = par.text.split(': ')[1]
        new_doc = Document()
    
    get_para_data(new_doc, par)  # Add text to existing doc
